"""
Word 导出 Skill
==============
将合并后的 Markdown 转为 DOCX 文件并保存到会话导出目录。
支持将 Mermaid 代码块渲染为 PNG 图片，嵌入文档末尾附件章节。
"""

import os
import re
import logging
from typing import List, Tuple

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from skills.export_utils import get_export_dir, build_combined_markdown

logger = logging.getLogger("agent_skills.export_docx")


def export_to_docx(
    session_id: str,
    project_name: str,
    contents: List[Tuple[str, str]],
    filename_suffix: str = "_需求与详设",
) -> Tuple[bool, str, str]:
    """
    将多段 (标题, Markdown内容) 合并并导出为 Word 文档。
    自动检测 Mermaid 代码块并渲染为 PNG 图片，放入文档末尾附件章节。

    Returns:
        (success, file_path, error_message)
    """
    export_dir = get_export_dir(session_id)
    filename = f"{project_name or 'document'}{filename_suffix}.docx"
    filename = _sanitize_filename(filename)
    file_path = os.path.join(export_dir, filename)

    try:
        combined_md = build_combined_markdown(contents, project_name)

        # Mermaid 预处理：提取代码块并渲染为 PNG
        mermaid_images: list[tuple[str, str]] = []
        try:
            from skills.mermaid_renderer import extract_and_render_mermaid
            img_dir = os.path.join(export_dir, "images")
            combined_md, mermaid_images = extract_and_render_mermaid(combined_md, img_dir)
        except Exception as mermaid_exc:
            logger.warning("Mermaid 预处理失败（保留原始代码块）: %s", mermaid_exc)

        doc = Document()
        doc.styles["Normal"].font.size = Pt(10.5)
        doc.styles["Normal"].font.name = "宋体"

        lines = combined_md.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if not line:
                i += 1
                continue

            # 标题
            if line.startswith("# "):
                p = doc.add_paragraph(line[2:].strip())
                p.style = "Heading 1"
                p.paragraph_format.space_after = Pt(12)
                i += 1
                continue
            elif line.startswith("## "):
                p = doc.add_paragraph(line[3:].strip())
                p.style = "Heading 2"
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                i += 1
                continue
            elif line.startswith("### "):
                p = doc.add_paragraph(line[4:].strip())
                p.style = "Heading 3"
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                i += 1
                continue

            # Markdown 表格：| 表头 | 表头 |
            if line.strip().startswith("|") and "|" in line[1:]:
                i = _add_md_table(doc, lines, i)
                continue

            # 普通段落
            _add_md_paragraph(doc, line)
            i += 1

        # 在文档末尾添加附件章节（Mermaid 渲染的 PNG 图片）
        if mermaid_images:
            doc.add_page_break()
            p = doc.add_paragraph("附件")
            p.style = "Heading 1"
            p.paragraph_format.space_after = Pt(12)

            for img_title, img_path in mermaid_images:
                if os.path.isfile(img_path):
                    p_label = doc.add_paragraph(f"图: {img_title}")
                    p_label.style = "Heading 3"
                    p_label.paragraph_format.space_before = Pt(12)
                    p_label.paragraph_format.space_after = Pt(6)
                    try:
                        doc.add_picture(img_path, width=Inches(5.5))
                    except Exception as img_exc:
                        logger.warning("Word 插入图片失败 (%s): %s", img_path, img_exc)
                        doc.add_paragraph(f"[图片加载失败: {img_title}]")
                    doc.add_paragraph("")

        doc.save(file_path)
        logger.info("DOCX 已保存: %s", file_path)
        return True, file_path, ""
    except Exception as exc:
        logger.exception("DOCX 导出失败: %s", exc)
        return False, "", str(exc)


def _add_md_paragraph(doc: Document, line: str) -> None:
    """将一行 Markdown 转为 Word 段落（简单处理 * ** 等）。"""
    text = line.strip()
    if not text:
        return
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = text.replace("\\n", "\n")
    doc.add_paragraph(text)


def _add_md_table(doc: Document, lines: list[str], start: int) -> int:
    """
    将从 start 行开始的一段 Markdown 表格渲染为 Word 表格。

    Markdown 结构示例：
    | 列1 | 列2 |
    | --- | --- |
    | a   | b   |
    """
    rows: list[str] = []
    i = start
    while i < len(lines):
        raw = lines[i].strip()
        if not raw or not raw.startswith("|") or "|" not in raw[1:]:
            break
        rows.append(raw)
        i += 1

    if len(rows) < 2:
        # 非标准表格，退化为普通段落
        for r in rows:
            doc.add_paragraph(r)
        return i

    header_cells = [c.strip() for c in rows[0].strip("|").split("|")]
    data_rows = rows[2:] if len(rows) >= 3 else []

    cols = len(header_cells)
    table = doc.add_table(rows=len(data_rows) + 1, cols=cols)
    table.style = "Table Grid"

    # 表头
    for idx, cell_text in enumerate(header_cells):
        if idx >= cols:
            break
        cell = table.cell(0, idx)
        cell.text = cell_text

    # 数据行
    for row_idx, row in enumerate(data_rows, start=1):
        cells = [c.strip() for c in row.strip("|").split("|")]
        for col_idx in range(min(cols, len(cells))):
            cell = table.cell(row_idx, col_idx)
            cell.text = cells[col_idx]

    return i


def _sanitize_filename(name: str) -> str:
    for c in ["/", "\\", ":", "*", "?", '"', "<", ">", "|"]:
        name = name.replace(c, "_")
    return name[:200] if len(name) > 200 else name
