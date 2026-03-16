"""
文件内容提取器
=============
从用户上传的文件中提取纯文本内容，支持多种常见格式。

支持的文件类型：
- 纯文本：.txt, .md, .csv, .json, .xml, .yaml, .yml, .log
- Word 文档：.docx
- PDF 文档：.pdf
- 其他文本：尝试以 UTF-8 读取
"""

import os
import io
import json
import logging
from typing import Optional

logger = logging.getLogger("agent_skills.file_reader")

# 直接以文本模式读取的扩展名集合
_TEXT_EXTENSIONS = {
    ".txt", ".md", ".csv", ".json", ".xml", ".yaml", ".yml",
    ".log", ".html", ".htm", ".css", ".js", ".ts", ".py",
    ".java", ".sql", ".sh", ".bat", ".ini", ".cfg", ".conf",
    ".properties", ".toml", ".rst", ".tex",
}

# 上传文件大小上限（10 MB）
MAX_FILE_SIZE = 10 * 1024 * 1024

# 提取内容的字符上限，防止 LLM 上下文过长
MAX_CONTENT_CHARS = 80_000


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, Optional[str]]:
    """
    从文件字节内容中提取纯文本。

    Args:
        file_bytes: 文件二进制内容
        filename: 原始文件名（用于推断格式）

    Returns:
        (extracted_text, error_message)
        - 成功时 error_message 为 None
        - 失败时 extracted_text 为空字符串
    """
    if len(file_bytes) > MAX_FILE_SIZE:
        return "", f"文件大小超过限制（最大 {MAX_FILE_SIZE // 1024 // 1024} MB）"

    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == ".docx":
            text = _extract_docx(file_bytes)
        elif ext == ".pdf":
            text = _extract_pdf(file_bytes)
        elif ext in _TEXT_EXTENSIONS:
            text = _extract_plain_text(file_bytes)
        else:
            # 尝试以文本方式读取
            text = _extract_plain_text(file_bytes)

        if not text or not text.strip():
            return "", "文件内容为空或无法提取文本"

        # 截断过长内容
        if len(text) > MAX_CONTENT_CHARS:
            text = text[:MAX_CONTENT_CHARS] + f"\n\n... （内容已截断，原文共 {len(text)} 字符）"

        logger.info("成功提取文件 '%s' 的文本内容，共 %d 字符", filename, len(text))
        return text, None

    except Exception as exc:
        logger.exception("提取文件 '%s' 内容失败: %s", filename, exc)
        return "", f"文件内容提取失败: {exc}"


def _extract_plain_text(file_bytes: bytes) -> str:
    """提取纯文本文件内容，尝试 UTF-8 和 GBK 编码。"""
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _extract_docx(file_bytes: bytes) -> str:
    """
    提取 Word (.docx) 文档中的文本内容。
    按段落提取，保留段落分隔。
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx 库来读取 Word 文档: pip install python-docx")

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 保留标题层级信息
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                try:
                    level_num = int(level)
                except ValueError:
                    level_num = 1
                paragraphs.append("#" * level_num + " " + text)
            else:
                paragraphs.append(text)

    # 提取表格内容
    for table in doc.tables:
        table_lines = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                table_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        if table_lines:
            paragraphs.append("\n".join(table_lines))

    return "\n\n".join(paragraphs)


def _extract_pdf(file_bytes: bytes) -> str:
    """
    提取 PDF 文档中的文本内容。
    优先尝试 PyPDF2，回退到 reportlab 的简单提取。
    """
    # 尝试 PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        if pages:
            return "\n\n".join(pages)
    except ImportError:
        logger.debug("PyPDF2 未安装，尝试其他方式提取 PDF")
    except Exception as exc:
        logger.warning("PyPDF2 提取 PDF 失败: %s", exc)

    # 尝试 pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            if pages:
                return "\n\n".join(pages)
    except ImportError:
        logger.debug("pdfplumber 未安装")
    except Exception as exc:
        logger.warning("pdfplumber 提取 PDF 失败: %s", exc)

    raise RuntimeError("无法提取 PDF 内容，请安装 PyPDF2: pip install PyPDF2")


def get_file_type_label(filename: str) -> str:
    """根据文件扩展名返回用户友好的文件类型标签。"""
    ext = os.path.splitext(filename)[1].lower()
    labels = {
        ".txt": "文本文件",
        ".md": "Markdown 文件",
        ".csv": "CSV 数据文件",
        ".json": "JSON 文件",
        ".xml": "XML 文件",
        ".yaml": "YAML 文件",
        ".yml": "YAML 文件",
        ".docx": "Word 文档",
        ".pdf": "PDF 文档",
        ".html": "HTML 文件",
        ".py": "Python 文件",
        ".java": "Java 文件",
        ".js": "JavaScript 文件",
        ".ts": "TypeScript 文件",
        ".sql": "SQL 文件",
    }
    return labels.get(ext, "文件")
